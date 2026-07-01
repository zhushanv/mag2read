#!/usr/bin/env python3
"""
从 clean/document.json 导出多种格式：EPUB、DOCX、TXT、HTML。

用法:
  conda run -n industrial-cv python backend/scripts/export_document.py DDCPC_sample --formats epub,docx,txt,html
  conda run -n industrial-cv python backend/scripts/export_document.py DDCPC_sample --formats epub  (默认)
"""

from __future__ import annotations

import argparse
import base64
import copy
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from ebooklib import epub
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template

from backend.app.modules.document_build import build_chapters, build_document


PROJECT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_STORAGE_ROOT = PROJECT_ROOT / "backend" / "storage" / "tasks"


def project_relative(path: Path) -> str:
    return str(path.relative_to(PROJECT_ROOT)) if path.is_relative_to(PROJECT_ROOT) else str(path)


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def resolve_task_dir(task: str, storage_root: Path) -> Path:
    p = Path(task).expanduser()
    if p.exists():
        return p.resolve()
    return (storage_root / task).resolve()


# ---------------------------------------------------------------
# 读取清洗后文档
# ---------------------------------------------------------------

def load_document(task_dir: Path) -> dict[str, Any]:
    return build_document(task_dir)


# ---------------------------------------------------------------
# 通用文本工具
# ---------------------------------------------------------------

def strip_title_newlines(text: str) -> str:
    """将标题中的换行替换为空格，使标题在一行显示。"""
    return re.sub(r"\s*\n\s*", " ", text).strip()


def iter_blocks(document: dict[str, Any]):
    """遍历文档中所有页面和 block。"""
    for page in document.get("pages", []):
        page_no = page["page_no"]
        for block in page.get("blocks", []):
            yield page_no, block


def media_path_for_block(block: dict[str, Any]) -> Path | None:
    value = str(block.get("media_path") or "").strip()
    if not value:
        return None
    path = Path(value)
    if not path.is_absolute():
        path = PROJECT_ROOT / path
    return path.resolve()


def block_has_media(block: dict[str, Any]) -> bool:
    role = str(block.get("role") or block.get("type") or "").lower()
    return bool(block.get("media_path")) or bool(block.get("is_graphical")) or role in {"figure", "image", "table", "formula"}


def media_caption(block: dict[str, Any]) -> str:
    return str(block.get("text") or "").strip()


def relative_media_href(block: dict[str, Any], base_dir: Path) -> str | None:
    path = media_path_for_block(block)
    if path is None or not path.exists():
        return None
    return os.path.relpath(path, base_dir)


def embedded_media_data_uri(block: dict[str, Any]) -> str | None:
    path = media_path_for_block(block)
    if path is None or not path.exists():
        return None
    suffix = path.suffix.lower()
    mime_type = "image/png"
    if suffix in {".jpg", ".jpeg"}:
        mime_type = "image/jpeg"
    elif suffix == ".webp":
        mime_type = "image/webp"
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def filter_document_media(document: dict[str, Any], include_media: bool) -> dict[str, Any]:
    if include_media:
        return document
    filtered = copy.deepcopy(document)
    pages: list[dict[str, Any]] = []
    for page in filtered.get("pages", []):
        next_page = dict(page)
        next_page["blocks"] = [
            block for block in page.get("blocks", [])
            if not block_has_media(block)
        ]
        pages.append(next_page)
    filtered["pages"] = pages
    filtered["chapters"] = build_chapters(pages)
    if isinstance(filtered.get("blocks"), list):
        filtered["blocks"] = [
            block for block in filtered["blocks"]
            if not block_has_media(block)
        ]
    filtered["include_media"] = False
    return filtered


# ---------------------------------------------------------------
# EPUB 导出
# ---------------------------------------------------------------

def export_epub(document: dict[str, Any], output_path: Path) -> str:
    title = strip_title_newlines(document.get("title", "Untitled"))
    author = document.get("author", "")
    language = document.get("language", "zh-CN")

    book = epub.EpubBook()
    book.set_identifier(f"mag2read-{document.get('task_id', 'unknown')}")
    book.set_title(title)
    book.set_language(language)
    if author:
        book.add_author(author)

    # 添加样式
    style = """
    body { font-family: serif; line-height: 1.8; margin: 1em; }
    h1 { font-size: 1.6em; margin-top: 1.5em; margin-bottom: 0.5em; }
    h2 { font-size: 1.3em; margin-top: 1.2em; margin-bottom: 0.4em; }
    h3 { font-size: 1.1em; margin-top: 1em; margin-bottom: 0.3em; }
    p { text-indent: 2em; margin: 0.5em 0; }
    p.no-indent { text-indent: 0; }
    blockquote {
        margin: 0.5em 1em; padding: 0.5em 1em;
        border-left: 3px solid #ccc; color: #555;
        font-style: italic;
    }
    .caption { font-size: 0.9em; color: #666; margin: 0.3em 0; }
    .sidebar { background: #f5f5f5; padding: 0.5em 1em; margin: 0.5em 0; border-radius: 4px; }
    figure.media { margin: 1em 0; text-align: center; }
    figure.media img { max-width: 100%; height: auto; }
    figure.media figcaption { font-size: 0.9em; color: #666; margin-top: 0.4em; }
    """
    css = epub.EpubItem(
        uid="style",
        file_name="style.css",
        media_type="text/css",
        content=style,
    )
    book.add_item(css)
    epub_media = build_epub_media_items(document)
    for item in epub_media.values():
        book.add_item(item["item"])

    chapters_epub: list[epub.EpubHtml] = []
    spine = ["nav"]

    # 按 chapter 组织（如果有 chapters），否则按 page
    chapters_data = document.get("chapters", [])

    if chapters_data:
        for ch in chapters_data:
            ch_title = strip_title_newlines(ch.get("title", "Untitled"))
            content_parts: list[str] = [
                f'<h1>{_escape_html(ch_title)}</h1>'
            ]
            for block in ch.get("blocks", []):
                content_parts.append(_block_to_html(block, epub_media=epub_media))
            html_content = _wrap_html(ch_title, "\n".join(content_parts), style_ref="style.css")
            chapter = epub.EpubHtml(
                title=ch_title,
                file_name=f"chap_{ch['chapter_id']}.xhtml",
                lang=language,
            )
            chapter.content = html_content
            chapter.add_item(css)
            book.add_item(chapter)
            chapters_epub.append(chapter)
            spine.append(chapter)
    else:
        # 无章节结构，按 page 组织
        for page in document.get("pages", []):
            page_no = page["page_no"]
            content_parts: list[str] = []
            for block in page.get("blocks", []):
                content_parts.append(_block_to_html(block, epub_media=epub_media))
            html_content = _wrap_html(
                f"Page {page_no}",
                "\n".join(content_parts),
                style_ref="style.css",
            )
            chapter = epub.EpubHtml(
                title=f"Page {page_no}",
                file_name=f"page_{page_no:03d}.xhtml",
                lang=language,
            )
            chapter.content = html_content
            chapter.add_item(css)
            book.add_item(chapter)
            chapters_epub.append(chapter)
            spine.append(chapter)

    book.toc = chapters_epub
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine

    epub.write_epub(str(output_path), book)
    return str(output_path)


def build_epub_media_items(document: dict[str, Any]) -> dict[str, dict[str, Any]]:
    media: dict[str, dict[str, Any]] = {}
    used_names: set[str] = set()
    for _, block in iter_blocks(document):
        path = media_path_for_block(block)
        if path is None or not path.exists():
            continue
        name = path.name
        if name in used_names:
            name = f"{path.stem}_{len(used_names) + 1}{path.suffix}"
        used_names.add(name)
        href = f"images/{name}"
        media[str(path)] = {
            "href": href,
            "item": epub.EpubItem(
                uid=f"media_{len(media) + 1}",
                file_name=href,
                media_type="image/png",
                content=path.read_bytes(),
            ),
        }
    return media


def _block_to_html(block: dict[str, Any], epub_media: dict[str, dict[str, Any]] | None = None) -> str:
    block_type = block.get("type", "paragraph")
    text = _escape_html(str(block.get("text") or ""))
    is_para_break = block.get("_paragraph_break", False)
    if block_has_media(block):
        path = media_path_for_block(block)
        item = epub_media.get(str(path)) if epub_media and path is not None else None
        if item:
            caption = f"<figcaption>{text}</figcaption>" if text else ""
            return f'<figure class="media"><img src="{item["href"]}" alt="{text or block_type}"/>{caption}</figure>'
        if text:
            return f'<blockquote class="caption">{text}</blockquote>'
        return '<p class="no-indent">[图片缺失]</p>'

    if block_type == "heading":
        level = min(int(block.get("level", 2)), 6)
        return f"<h{level}>{text}</h{level}>"
    if block_type == "caption":
        return f'<blockquote class="caption">{text}</blockquote>'
    if block_type == "sidebar":
        return f'<div class="sidebar"><strong>补充：</strong>{text}</div>'
    if block_type == "note":
        return f'<div class="sidebar"><strong>注：</strong>{text}</div>'
    # paragraph
    indent_class = "" if is_para_break else ' class="no-indent"'
    return f"<p{indent_class}>{text}</p>"


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _wrap_html(title: str, body: str, style_ref: str = "style.css") -> str:
    return f"""<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><meta charset="utf-8"/><title>{_escape_html(title)}</title>
<link rel="stylesheet" type="text/css" href="{style_ref}"/></head>
<body>{body}</body>
</html>"""


# ---------------------------------------------------------------
# DOCX 导出
# ---------------------------------------------------------------

def set_run_font(run, font_name='SimSun', size=Pt(11), bold=False, italic=False, color=None):
    """设置 run 的中英文字体，确保中文字体正确生效。"""
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # 设置东亚字体（中文）
    from docx.oxml.ns import qn
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_indent(p, first_line_cm=0.74):
    """设置段落首行缩进（两字符约0.74cm）"""
    p.paragraph_format.first_line_indent = Cm(first_line_cm)


def export_docx(document: dict[str, Any], output_path: Path) -> str:
    doc = Document()

    # 页面设置 — A4 尺寸，适合中文文档
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # 默认样式
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    # Normal 样式设置东亚字体
    from docx.oxml.ns import qn
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = style.element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 修改内置标题样式 — 使用黑体
    for level in range(1, 4):
        hs_key = f'Heading {level}'
        if hs_key in doc.styles:
            hs = doc.styles[hs_key]
            hs.font.name = 'SimHei'
            hs.font.bold = True
            hs.font.color.rgb = None
            hrpr = hs.element.get_or_add_rPr()
            hrFonts = hrpr.find(qn('w:rFonts'))
            if hrFonts is None:
                hrFonts = hs.element.makeelement(qn('w:rFonts'), {})
                hrpr.insert(0, hrFonts)
            hrFonts.set(qn('w:eastAsia'), 'SimHei')
            if level == 1:
                hs.font.size = Pt(16)
            elif level == 2:
                hs.font.size = Pt(14)
            else:
                hs.font.size = Pt(13)

    # 正文
    chapters_data = document.get("chapters", [])
    if chapters_data:
        for ch in chapters_data:
            ch_title = strip_title_newlines(ch.get("title", ""))
            if ch_title:
                heading = doc.add_heading(ch_title, level=1)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)
            for block in ch.get("blocks", []):
                _add_block_to_docx(doc, block)
    else:
        for page in document.get("pages", []):
            for block in page.get("blocks", []):
                _add_block_to_docx(doc, block)

    doc.save(str(output_path))
    return str(output_path)


def _add_block_to_docx(doc: Document, block: dict[str, Any]) -> None:
    block_type = block.get("type", "paragraph")
    text = str(block.get("text") or "")
    is_para_break = block.get("_paragraph_break", False)
    if block_has_media(block):
        path = media_path_for_block(block)
        if path is not None and path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run()
            run.add_picture(str(path), width=Cm(14))
            if text:
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(text)
                set_run_font(caption_run, 'SimSun', Pt(9), italic=True)
                caption.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph("[图片缺失]")
            p.paragraph_format.left_indent = Cm(0.5)
        return

    if block_type == "heading":
        level = min(int(block.get("level", 2)), 6)
        p = doc.add_heading(text, level=level)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    elif block_type == "caption":
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, 'SimSun', Pt(10), italic=True)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(2)
    elif block_type in ("sidebar", "note"):
        prefix = "【补充】" if block_type == "sidebar" else "【注】"
        p = doc.add_paragraph()
        run = p.add_run(prefix + text)
        set_run_font(run, 'SimSun', Pt(10.5))
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
    else:
        # 正文段落 — 首行缩进两字符，不额外插入"正文"文字
        p = doc.add_paragraph(text)
        set_run_font(p.runs[0], 'SimSun', Pt(12)) if p.runs else None
        set_paragraph_indent(p, 0.74)
        p.paragraph_format.space_after = Pt(1)
        if is_para_break:
            pass  # 首行缩进已默认设置


# ---------------------------------------------------------------
# TXT 导出
# ---------------------------------------------------------------

def export_txt(document: dict[str, Any], output_path: Path) -> str:
    lines: list[str] = []

    chapters_data = document.get("chapters", [])
    if chapters_data:
        for ch in chapters_data:
            ch_title = strip_title_newlines(ch.get("title", ""))
            if ch_title:
                lines.append(ch_title)
                lines.append("-" * len(ch_title))
                lines.append("")
            for block in ch.get("blocks", []):
                _add_block_to_txt(lines, block)
    else:
        for page in document.get("pages", []):
            lines.append(f"--- Page {page['page_no']} ---")
            lines.append("")
            for block in page.get("blocks", []):
                _add_block_to_txt(lines, block)

    output_path.write_text("\n".join(lines), encoding="utf-8")
    return str(output_path)


def _add_block_to_txt(lines: list[str], block: dict[str, Any]) -> None:
    text = str(block.get("text") or "")
    block_type = block.get("type", "paragraph")
    if block_has_media(block):
        path = media_path_for_block(block)
        if path is not None and path.exists():
            lines.append(f"[图片：{path.name}]")
        else:
            lines.append("[图片缺失]")
        if text:
            lines.append(text)
        lines.append("")
        return

    if block_type == "heading":
        level = int(block.get("level", 2))
        marker = "#" * level
        lines.append(f"{marker} {text}")
    elif block_type == "caption":
        lines.append(f"> {text}")
    elif block_type == "sidebar":
        lines.append(f"> 补充：{text}")
    elif block_type == "note":
        lines.append(f"> 注：{text}")
    else:
        if block.get("_paragraph_break"):
            lines.append("")
        lines.append(text)
    lines.append("")


# ---------------------------------------------------------------
# HTML 导出（独立文件，便于浏览器预览）
# ---------------------------------------------------------------

HTML_TEMPLATE = Template("""<!DOCTYPE html>
<html lang="{{ lang }}">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{ title }}</title>
<style>
  body { font-family: 'Noto Serif', 'Source Han Serif SC', Georgia, serif;
         max-width: 720px; margin: 2em auto; padding: 0 1em;
         line-height: 1.9; color: #222; }
  h1 { font-size: 1.5em; margin-top: 1.2em; margin-bottom: 0.6em;
       border-bottom: 2px solid #ddd; padding-bottom: 0.3em; }
  h2 { font-size: 1.25em; margin-top: 1em; margin-bottom: 0.4em; }
  h3 { font-size: 1.1em; margin-top: 0.8em; margin-bottom: 0.3em; }
  p { text-indent: 2em; margin: 0.4em 0; }
  p.no-indent { text-indent: 0; }
  blockquote { margin: 0.5em 1em; padding: 0.3em 1em;
               border-left: 3px solid #aaa; color: #555; }
  blockquote.caption { font-style: italic; font-size: 0.9em; }
  .sidebar { background: #f5f5f5; padding: 0.5em 1em;
             border-radius: 4px; margin: 0.5em 0; font-size: 0.95em; }
  figure.media { margin: 1em 0; text-align: center; }
  figure.media img { max-width: 100%; height: auto; border-radius: 4px; }
  figure.media figcaption { margin-top: 0.4em; color: #666; font-size: 0.9em; }
  .page-marker { text-align: center; color: #999; font-size: 0.85em;
                 margin: 1.5em 0; border-top: 1px dashed #ddd;
                 padding-top: 0.5em; }
  hr { border: none; border-top: 1px solid #eee; margin: 1.5em 0; }
</style>
</head>
<body>
  {% if author %}<p style="text-align:center;color:#666;">{{ author }}</p>{% endif %}
  {% if summary %}<blockquote><p>{{ summary }}</p></blockquote>{% endif %}
  {% for page in pages %}
  <div class="page-marker">— Page {{ page.page_no }} —</div>
  {% for block in page.blocks %}
    {{ block_html(block) }}
  {% endfor %}
  {% endfor %}
</body>
</html>""")


HTML_BLOCK_MAP = {
    "heading": lambda b: f"<h{min(b.get('level', 2), 6)}>{_escape_html(b['text'])}</h{min(b.get('level', 2), 6)}>",
    "caption": lambda b: f'<blockquote class="caption">{_escape_html(b["text"])}</blockquote>',
    "sidebar": lambda b: f'<div class="sidebar"><strong>补充：</strong>{_escape_html(b["text"])}</div>',
    "note": lambda b: f'<div class="sidebar"><strong>注：</strong>{_escape_html(b["text"])}</div>',
}


def _html_for_block(block: dict[str, Any], base_dir: Path | None = None, embed_media: bool = True) -> str:
    block_type = block.get("type", "paragraph")
    if block_has_media(block):
        href = embedded_media_data_uri(block) if embed_media else relative_media_href(block, base_dir or PROJECT_ROOT)
        caption = _escape_html(media_caption(block))
        if href:
            caption_html = f"<figcaption>{caption}</figcaption>" if caption else ""
            return f'<figure class="media"><img src="{_escape_html(href)}" alt="{caption or block_type}">{caption_html}</figure>'
        if caption:
            return f'<blockquote class="caption">{caption}</blockquote>'
        return '<p class="no-indent">[图片缺失]</p>'
    renderer = HTML_BLOCK_MAP.get(block_type)
    if renderer:
        return renderer(block)
    text = _escape_html(str(block.get("text") or ""))
    cls = "" if block.get("_paragraph_break") else ' class="no-indent"'
    return f"<p{cls}>{text}</p>"


def export_html(document: dict[str, Any], output_path: Path) -> str:
    title = strip_title_newlines(document.get("title", "Untitled"))
    author = document.get("author", "")
    summary = document.get("summary", "")
    pages = document.get("pages", [])

    html = HTML_TEMPLATE.render(
        title=title,
        author=author,
        summary=summary,
        lang=document.get("language", "zh-CN"),
        pages=pages,
        block_html=lambda block: _html_for_block(block, output_path.parent, embed_media=False),
    )
    output_path.write_text(html, encoding="utf-8")
    return str(output_path)


# ---------------------------------------------------------------
# 统一导出入口
# ---------------------------------------------------------------
def export_markdown(document: dict[str, Any], output_path: Path) -> str:
    lines: list[str] = []

    author = document.get("author", "")
    if author:
        lines.append(f"> 作者：{author}")
        lines.append("")

    chapters_data = document.get("chapters", [])
    if chapters_data:
        for ch in chapters_data:
            ch_title = strip_title_newlines(ch.get("title", ""))
            if ch_title:
                lines.append(f"## {ch_title}")
                lines.append("")
            for block in ch.get("blocks", []):
                _add_block_to_md(lines, block, output_path.parent)
    else:
        for page in document.get("pages", []):
            lines.append(f"---")
            lines.append(f"<small>第 {page['page_no']} 页</small>")
            lines.append("")
            for block in page.get("blocks", []):
                _add_block_to_md(lines, block, output_path.parent)

    output_path.write_text("\n".join(lines), encoding="utf-8")
    return str(output_path)


def _add_block_to_md(lines: list[str], block: dict[str, Any], base_dir: Path) -> None:
    text = str(block.get("text") or "")
    block_type = block.get("type", "paragraph")
    if block_has_media(block):
        href = relative_media_href(block, base_dir)
        if href:
            alt = text or block.get("role") or block_type or "图片"
            lines.append(f"![{alt}]({href})")
        else:
            lines.append("[图片缺失]")
        if text:
            lines.append("")
            lines.append(f"*{text}*")
        lines.append("")
        return

    if block_type == "heading":
        level = min(int(block.get("level", 2)), 6)
        lines.append(f"{'#' * level} {text}")
    elif block_type == "caption":
        lines.append(f"*{text}*")
    elif block_type == "sidebar":
        lines.append(f"> **补充：** {text}")
    elif block_type == "note":
        lines.append(f"> **注：** {text}")
    else:
        if block.get("_paragraph_break"):
            lines.append("")
        lines.append(text)
    lines.append("")


# ---------------------------------------------------------------
# 统一导出入口
# ---------------------------------------------------------------

FORMAT_EXPORTERS = {
    "epub": ("EPUB 电子书", export_epub, ".epub"),
    "docx": ("Word 文档", export_docx, ".docx"),
    "txt": ("纯文本", export_txt, ".txt"),
    "html": ("HTML 网页", export_html, ".html"),
    "markdown": ("Markdown", export_markdown, ".md"),
}


def run_export(task_dir: Path, formats: list[str], include_media: bool = True) -> dict[str, str]:
    document = filter_document_media(load_document(task_dir), include_media)
    output_dir = task_dir / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_variant = "with-media" if include_media else "text-only"

    results: dict[str, str] = {}

    for fmt in formats:
        if fmt not in FORMAT_EXPORTERS:
            print(f"  跳过未知格式: {fmt}")
            continue
        label, exporter, suffix = FORMAT_EXPORTERS[fmt]
        output_path = output_dir / f"{task_dir.name}-{output_variant}{suffix}"
        try:
            out = exporter(document, output_path)
            results[fmt] = project_relative(Path(out))
            print(f"  {label:10s} → {project_relative(Path(out))}")
        except Exception as exc:
            print(f"  {label:10s} ✗ 失败: {exc}")

    return results


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="从 clean/document.json 导出多种格式")
    parser.add_argument("task", help="Task ID 或完整路径")
    parser.add_argument(
        "--formats", "-f",
        default="epub",
        help="导出格式，逗号分隔。可选: epub,docx,txt,html。默认: epub",
    )
    parser.add_argument(
        "--storage-root",
        default=str(DEFAULT_STORAGE_ROOT),
    )
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    task_dir = resolve_task_dir(args.task, Path(args.storage_root).expanduser().resolve())
    formats = [fmt.strip().lower() for fmt in args.formats.split(",") if fmt.strip()]

    print(f"导出任务: {task_dir.name}")
    print(f"目标格式: {', '.join(formats)}")
    results = run_export(task_dir, formats)

    if not results:
        print("没有成功导出的文件。")
        return 1

    print(f"\n导出完成，共 {len(results)} 个文件。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
